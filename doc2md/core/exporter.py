"""Export markdown content to multiple formats (.md, .txt, .docx)."""

from pathlib import Path
from typing import Literal


def export_markdown(content: str, output_path: Path, format_type: Literal["md", "txt", "docx"] = "md") -> tuple[bool, str]:
    """
    Export markdown content to specified format.

    Returns: (success: bool, message: str)
    """
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if format_type == "md":
            output_path.write_text(content, encoding="utf-8")
            return True, f"Exported to {output_path.name}"

        elif format_type == "txt":
            output_path.write_text(content, encoding="utf-8")
            return True, f"Exported to {output_path.name}"

        elif format_type == "docx":
            try:
                from docx import Document
                from docx.shared import Pt, RGBColor
                from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            except ImportError:
                return False, "python-docx not installed. Install via: pip install 'doc2md[gui]'"

            doc = Document()

            for line in content.split("\n"):
                line = line.rstrip()
                if not line:
                    doc.add_paragraph()
                    continue

                if line.startswith("# "):
                    p = doc.add_heading(line[2:], level=1)
                elif line.startswith("## "):
                    p = doc.add_heading(line[3:], level=2)
                elif line.startswith("### "):
                    p = doc.add_heading(line[4:], level=3)
                elif line.startswith("#### "):
                    p = doc.add_heading(line[5:], level=4)
                elif line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("  "):
                    p = doc.add_paragraph(line.lstrip(), style="List Bullet 2")
                else:
                    p = doc.add_paragraph(line)

                if "**" in line:
                    for run in p.runs:
                        if "**" in run.text:
                            run.bold = True

            doc.save(str(output_path))
            return True, f"Exported to {output_path.name}"

        else:
            return False, f"Unknown format: {format_type}"

    except Exception as exc:
        return False, f"Export error: {exc}"
